from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


BASE_DIR = Path(__file__).resolve().parents[1]
OUT_DIR = BASE_DIR / "docs" / "manuales"
OUT_DIR.mkdir(parents=True, exist_ok=True)


def add_cover(doc: Document, role_title: str) -> None:
    p = doc.add_paragraph("MANUAL DE USUARIO")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(20)

    p2 = doc.add_paragraph(role_title)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].bold = True
    p2.runs[0].font.size = Pt(16)

    doc.add_paragraph("")
    fields = [
        "Aplicativo: LockerBeef",
        "Version del documento: 1.0",
        "Version del sistema: [completar]",
        "Fecha: [dd/mm/aaaa]",
        "Elaborado por: [completar]",
        "Aprobado por: [completar]",
    ]
    for field in fields:
        doc.add_paragraph(field)

    doc.add_page_break()


def add_section(doc: Document, title: str, body_lines: list[str]) -> None:
    doc.add_heading(title, level=1)
    for line in body_lines:
        if line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line[:3] in {"1. ", "2. ", "3. ", "4. ", "5. "}:
            doc.add_paragraph(line, style="List Number")
        else:
            doc.add_paragraph(line)


def build_admin() -> None:
    doc = Document()
    add_cover(doc, "Rol Administrador")

    add_section(
        doc,
        "Introduccion",
        [
            "Este manual orienta al usuario con rol Administrador en el uso del aplicativo LockerBeef para consulta y gestion operativa de lockers, dotaciones y personal.",
            "Incluye los pasos de acceso, navegacion, funciones principales, mensajes frecuentes y buenas practicas para una operacion segura y consistente.",
        ],
    )
    add_section(
        doc,
        "1. Objetivo",
        [
            "Describir de forma clara las actividades que puede ejecutar el rol Administrador dentro del sistema, asegurando un uso correcto de las funcionalidades habilitadas.",
        ],
    )
    add_section(
        doc,
        "2. Alcance",
        [
            "Aplica a usuarios con rol Administrador en ambiente de intranet, desde el ingreso al sistema hasta el cierre de sesion.",
        ],
    )
    add_section(
        doc,
        "3. Perfil del Rol Administrador",
        [
            "- Acceso al dashboard general.",
            "- Visualizacion de informacion por modulos y areas.",
            "- Gestion de usuarios (crear, editar, activar/inactivar segun politicas).",
            "- Registro y consulta operativa de lockers y dotacion.",
        ],
    )
    add_section(
        doc,
        "4. Requisitos Previos",
        [
            "- Equipo conectado a la red interna.",
            "- Navegador actualizado (Edge o Chrome).",
            "- Credenciales activas.",
            "- URL del aplicativo: [completar].",
        ],
    )
    add_section(
        doc,
        "5. Ingreso al Sistema",
        [
            "1. Abrir la URL del aplicativo.",
            "2. Ingresar correo y contrasena (o acceso integrado si aplica).",
            "3. Seleccionar area cuando el sistema lo solicite.",
            "4. Verificar carga del dashboard.",
        ],
    )
    add_section(
        doc,
        "6. Navegacion General",
        [
            "- Barra lateral: acceso a modulos.",
            "- Encabezado: nombre de usuario, rol, tiempo de sesion y cierre de sesion.",
            "- Panel principal: indicadores y resumen de estado.",
        ],
    )
    add_section(
        doc,
        "7. Funcionalidades Principales",
        [
            "7.1 Dashboard: revisar indicadores globales y estado operativo.",
            "7.2 Gestion de usuarios: administrar cuentas y roles permitidos.",
            "7.3 Operacion: registros de personal, lockers y dotacion.",
            "7.4 Consultas: historial, filtros por estado, codigo y area.",
        ],
    )
    add_section(
        doc,
        "8. Mensajes Frecuentes",
        [
            "- Email o contrasena incorrectos: validar credenciales.",
            "- Cuenta inactiva: contactar al administrador del sistema.",
            "- Sin datos: verificar filtros y area seleccionada.",
        ],
    )
    add_section(
        doc,
        "9. Buenas Practicas",
        [
            "- No compartir credenciales.",
            "- Cerrar sesion al finalizar actividades.",
            "- Validar area activa antes de registrar informacion.",
            "- Reportar inconsistencias al equipo de soporte.",
        ],
    )
    add_section(
        doc,
        "10. Soporte",
        [
            "Canal de soporte: [correo/telefono]",
            "Horario: [completar]",
            "Responsable funcional: [completar]",
        ],
    )
    add_section(
        doc,
        "11. Control de Cambios",
        [
            "Version | Fecha | Autor | Descripcion",
            "1.0 | [dd/mm/aaaa] | [nombre] | Emision inicial",
        ],
    )

    doc.save(OUT_DIR / "Manual_Usuario_Rol_Administrador_LockerBeef.docx")


def build_standard() -> None:
    doc = Document()
    add_cover(doc, "Rol Estandar (Solo Lectura)")

    add_section(
        doc,
        "Introduccion",
        [
            "Este manual describe el uso del aplicativo LockerBeef para usuarios con rol Estandar, cuyo acceso esta limitado a la visualizacion de informacion de su area asignada.",
            "Su proposito es facilitar la consulta operativa sin permitir modificaciones, garantizando control y trazabilidad de la informacion.",
        ],
    )
    add_section(
        doc,
        "1. Objetivo",
        [
            "Guiar al usuario Estandar en el acceso y consulta de informacion del sistema dentro de los limites de permisos definidos.",
        ],
    )
    add_section(
        doc,
        "2. Alcance",
        [
            "Aplica a usuarios con perfil de solo lectura en intranet, desde inicio de sesion hasta cierre de sesion.",
        ],
    )
    add_section(
        doc,
        "3. Perfil del Rol Estandar",
        [
            "- Visualiza unicamente informacion del area asignada.",
            "- No puede crear, editar ni eliminar registros.",
            "- No tiene acceso a gestion de usuarios.",
        ],
    )
    add_section(
        doc,
        "4. Requisitos Previos",
        [
            "- Conexion a red interna.",
            "- Navegador actualizado.",
            "- Credenciales activas.",
            "- Area asignada correctamente por Administrador.",
        ],
    )
    add_section(
        doc,
        "5. Ingreso al Sistema",
        [
            "1. Abrir la URL del aplicativo.",
            "2. Ingresar credenciales o acceso integrado cuando aplique.",
            "3. Verificar que la informacion corresponda al area asignada.",
        ],
    )
    add_section(
        doc,
        "6. Consultas Disponibles",
        [
            "- Visualizacion de estados de lockers.",
            "- Consulta de dotaciones y disponibilidad.",
            "- Revision de historicos permitidos.",
        ],
    )
    add_section(
        doc,
        "7. Restricciones",
        [
            "- No puede registrar informacion nueva.",
            "- No puede modificar datos existentes.",
            "- No puede eliminar registros.",
            "- No puede cambiar configuraciones o roles.",
        ],
    )
    add_section(
        doc,
        "8. Mensajes Frecuentes",
        [
            "- Sin permisos para esta accion: la funcion no esta habilitada para tu rol.",
            "- Cuenta inactiva: solicitar activacion al Administrador.",
            "- Sin resultados: verificar filtros de busqueda.",
        ],
    )
    add_section(
        doc,
        "9. Buenas Practicas",
        [
            "- Mantener confidencialidad de credenciales.",
            "- Cerrar sesion al terminar.",
            "- Reportar hallazgos o inconsistencias.",
        ],
    )
    add_section(
        doc,
        "10. Soporte",
        [
            "Canal de soporte: [correo/telefono]",
            "Horario: [completar]",
            "Responsable funcional: [completar]",
        ],
    )
    add_section(
        doc,
        "11. Control de Cambios",
        [
            "Version | Fecha | Autor | Descripcion",
            "1.0 | [dd/mm/aaaa] | [nombre] | Emision inicial",
        ],
    )

    doc.save(OUT_DIR / "Manual_Usuario_Rol_Estandar_LockerBeef.docx")


if __name__ == "__main__":
    build_admin()
    build_standard()
    print(OUT_DIR)
